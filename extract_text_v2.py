import docx
import sys

def get_docx_content(path):
    doc = docx.Document(path)
    content = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(f"PARA: {para.text}")
    
    # Process tables
    for table in doc.tables:
        content.append("TABLE START")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            content.append(" | ".join(row_data))
        content.append("TABLE END")
        
    return '\n'.join(content)

if __name__ == "__main__":
    path = "agentic_ai_linguistics_lecture.docx"
    text = get_docx_content(path)
    with open("lecture_content_full.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("Extracted full content to lecture_content_full.txt")
